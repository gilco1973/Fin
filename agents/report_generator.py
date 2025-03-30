from typing import Dict, List
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.charts.barcharts import VerticalBarChart
from langchain.agents import AgentExecutor
from langchain_openai import ChatOpenAI
from langchain.agents import create_openai_functions_agent
from langchain.tools import BaseTool
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.messages import SystemMessage, HumanMessage
import os
from dotenv import load_dotenv
from datetime import datetime
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

load_dotenv()

class ReportGeneratorTool(BaseTool):
    name: str = "report_generator"
    description: str = "Generates PDF reports from financial analysis results"
    
    def _format_currency(self, amount: float) -> str:
        """Format amount as currency."""
        return f"${amount:,.2f}"
    
    def _create_pie_chart(self, data: Dict[str, float], title: str) -> Drawing:
        """Create a pie chart using reportlab."""
        drawing = Drawing(400, 200)
        pie = Pie()
        pie.x = 150
        pie.y = 25
        pie.width = 200
        pie.height = 150
        
        # Convert data to lists for the pie chart
        pie.data = list(data.values())
        pie.labels = list(data.keys())
        
        # Add some visual styling
        pie.slices.strokeWidth = 0.5
        for i in range(len(data)):
            pie.slices[i].fillColor = colors.HexColor(f"#{hash(pie.labels[i]) % 0xFFFFFF:06x}")
        
        drawing.add(pie)
        return drawing
    
    def _run(self, analyses: List[Dict], output_file: str) -> None:
        """Generate a PDF report from financial analysis results."""
        try:
            # Create PDF document
            doc = SimpleDocTemplate(output_file, pagesize=letter)
            styles = getSampleStyleSheet()
            elements = []
            
            # Title
            title = Paragraph("Financial Analysis Report", styles["Title"])
            elements.append(title)
            elements.append(Spacer(1, 12))
            
            # Debug: Print analyses structure
            print("\nAnalyses structure:")
            for i, analysis in enumerate(analyses):
                print(f"\nAnalysis {i+1}:")
                print(f"Keys: {analysis.keys()}")
                if "metadata" in analysis:
                    print(f"Metadata: {analysis['metadata']}")
                if "output" in analysis:
                    print(f"Output keys: {analysis['output'].keys()}")
                    if "analysis" in analysis["output"]:
                        print(f"Analysis keys: {analysis['output']['analysis'].keys()}")
            
            # Source Files
            elements.append(Paragraph("Source Files", styles["Heading2"]))
            for analysis in analyses:
                if "metadata" in analysis and "file_name" in analysis["metadata"]:
                    file_name = analysis["metadata"]["file_name"]
                    elements.append(Paragraph(f"• {file_name}", styles["Normal"]))
            elements.append(Spacer(1, 12))
            
            # Statement Periods
            elements.append(Paragraph("Statement Periods", styles["Heading2"]))
            for analysis in analyses:
                if "output" in analysis and "analysis" in analysis["output"]:
                    data = analysis["output"]["analysis"]
                    if "statement_period" in data and "metadata" in analysis:
                        period = data["statement_period"]
                        file_name = analysis["metadata"]["file_name"]
                        start_date = period.get('start_date', '')
                        end_date = period.get('end_date', '')
                        
                        elements.append(Paragraph(f"• {file_name}:", styles["Heading3"]))
                        if start_date:
                            elements.append(Paragraph(f"  From: {start_date}", styles["Normal"]))
                        if end_date:
                            elements.append(Paragraph(f"  To: {end_date}", styles["Normal"]))
                        elements.append(Spacer(1, 6))
            elements.append(Spacer(1, 12))
            
            # Financial Analysis
            elements.append(Paragraph("Financial Analysis", styles["Heading2"]))
            
            # Separate credit card and bank statements
            credit_card_analyses = []
            bank_analyses = []
            
            for analysis in analyses:
                if "output" in analysis and "analysis" in analysis["output"]:
                    data = analysis["output"]["analysis"]
                    # Check if it's a credit card statement by looking for credit card specific fields
                    if "balance_info" in data and "previous_balance" in data["balance_info"]:
                        credit_card_analyses.append(analysis)
                    else:
                        bank_analyses.append(analysis)
            
            # Credit Card Statement Comparison
            if credit_card_analyses:
                elements.append(Paragraph("Credit Card Statement Comparison", styles["Heading3"]))
                elements.append(Spacer(1, 6))
                
                # Prepare credit card comparison data
                cc_comparison_data = [["Metric"]]
                cc_statement_data = []
                
                for analysis in credit_card_analyses:
                    data = analysis["output"]["analysis"]
                    period = data["statement_period"]
                    end_date = period.get("end_date", "Unknown")
                    cc_comparison_data[0].append(end_date)
                    
                    balance_info = data["balance_info"]
                    statement_metrics = {
                        "Purchases": float(balance_info.get("purchases", 0)),
                        "Payments": float(balance_info.get("payments", 0)),
                        "New Balance": float(balance_info.get("new_balance", 0)),
                        "Interest": float(balance_info.get("interest", 0)),
                        "Fees": float(balance_info.get("fees", 0)),
                        "Cash Advances": float(balance_info.get("cash_advances", 0))
                    }
                    cc_statement_data.append(statement_metrics)
                
                # Create rows for credit card metrics
                cc_metrics = ["Purchases", "Payments", "New Balance", "Interest", "Fees", "Cash Advances"]
                for metric in cc_metrics:
                    row = [metric]
                    for statement in cc_statement_data:
                        row.append(f"${statement.get(metric, 0):,.2f}")
                    cc_comparison_data.append(row)
                
                # Add month-over-month changes for credit cards
                if len(cc_statement_data) > 1:
                    for metric in cc_metrics:
                        row = [f"{metric} Change %"]
                        for i in range(len(cc_statement_data)):
                            if i == 0:
                                row.append("N/A")
                            else:
                                current = cc_statement_data[i].get(metric, 0)
                                previous = cc_statement_data[i-1].get(metric, 0)
                                if previous != 0:
                                    change = ((current - previous) / previous) * 100
                                    row.append(f"{'↑' if change > 0 else '↓'}{abs(change):.1f}%")
                                else:
                                    row.append("N/A")
                        cc_comparison_data.append(row)
                
                # Create and style the credit card comparison table
                cc_table = Table(cc_comparison_data)
                cc_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                    ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]))
                elements.append(cc_table)
                elements.append(Spacer(1, 12))
            
            # Bank Statement Comparison
            if bank_analyses:
                elements.append(Paragraph("Bank Statement Comparison", styles["Heading3"]))
                elements.append(Spacer(1, 6))
                
                # Prepare bank comparison data
                bank_comparison_data = [["Metric"]]
                bank_statement_data = []
                
                for analysis in bank_analyses:
                    data = analysis["output"]["analysis"]
                    period = data["statement_period"]
                    end_date = period.get("end_date", "Unknown")
                    bank_comparison_data[0].append(end_date)
                    
                    balance_info = data["balance_info"]
                    statement_metrics = {
                        "Opening Balance": float(balance_info.get("opening", 0)),
                        "Closing Balance": float(balance_info.get("closing", 0)),
                        "Net Change": float(balance_info.get("change", 0)),
                        "Deposits": float(balance_info.get("deposits", 0)),
                        "Withdrawals": float(balance_info.get("withdrawals", 0))
                    }
                    bank_statement_data.append(statement_metrics)
                
                # Create rows for bank metrics
                bank_metrics = ["Opening Balance", "Closing Balance", "Net Change", "Deposits", "Withdrawals"]
                for metric in bank_metrics:
                    row = [metric]
                    for statement in bank_statement_data:
                        row.append(f"${statement.get(metric, 0):,.2f}")
                    bank_comparison_data.append(row)
                
                # Add month-over-month changes for bank statements
                if len(bank_statement_data) > 1:
                    for metric in bank_metrics:
                        row = [f"{metric} Change %"]
                        for i in range(len(bank_statement_data)):
                            if i == 0:
                                row.append("N/A")
                            else:
                                current = bank_statement_data[i].get(metric, 0)
                                previous = bank_statement_data[i-1].get(metric, 0)
                                if previous != 0:
                                    change = ((current - previous) / previous) * 100
                                    row.append(f"{'↑' if change > 0 else '↓'}{abs(change):.1f}%")
                                else:
                                    row.append("N/A")
                        bank_comparison_data.append(row)
                
                # Create and style the bank comparison table
                bank_table = Table(bank_comparison_data)
                bank_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                    ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]))
                elements.append(bank_table)
                elements.append(Spacer(1, 12))
            
            # Add insights about the comparison
            elements.append(Paragraph("Key Insights:", styles["Heading3"]))
            
            # Calculate and add insights for credit cards
            if credit_card_analyses:
                elements.append(Paragraph("Credit Card Insights:", styles["Heading4"]))
                cc_insights = []
                
                # Analyze credit card trends
                if len(cc_statement_data) > 1:
                    # Spending pattern
                    total_purchases = sum(data.get("Purchases", 0) for data in cc_statement_data)
                    avg_purchases = total_purchases / len(cc_statement_data)
                    latest_purchases = cc_statement_data[0].get("Purchases", 0)
                    if latest_purchases > avg_purchases:
                        cc_insights.append(f"• Recent purchases (${latest_purchases:,.2f}) are above the {len(cc_statement_data)}-month average of ${avg_purchases:,.2f}")
                    else:
                        cc_insights.append(f"• Recent purchases (${latest_purchases:,.2f}) are below the {len(cc_statement_data)}-month average of ${avg_purchases:,.2f}")
                    
                    # Payment behavior
                    total_payments = sum(data.get("Payments", 0) for data in cc_statement_data)
                    payment_ratio = total_payments / total_purchases if total_purchases > 0 else 0
                    if payment_ratio > 1:
                        cc_insights.append(f"• Strong payment behavior: Payments exceed purchases by {((payment_ratio - 1) * 100):.1f}%")
                    elif payment_ratio > 0.9:
                        cc_insights.append("• Good payment behavior: Most purchases are being paid off")
                    else:
                        cc_insights.append(f"• Warning: Payments cover only {(payment_ratio * 100):.1f}% of purchases")
                    
                    # Interest and fees
                    total_interest = sum(data.get("Interest", 0) for data in cc_statement_data)
                    total_fees = sum(data.get("Fees", 0) for data in cc_statement_data)
                    if total_interest > 0 or total_fees > 0:
                        cc_insights.append(f"• Total interest and fees paid: ${(total_interest + total_fees):,.2f}")
                    
                    # Balance trend
                    first_balance = cc_statement_data[-1].get("New Balance", 0)
                    latest_balance = cc_statement_data[0].get("New Balance", 0)
                    balance_change = ((latest_balance - first_balance) / first_balance * 100) if first_balance != 0 else 0
                    if balance_change > 0:
                        cc_insights.append(f"• Balance has increased by {balance_change:.1f}% over the period")
                    else:
                        cc_insights.append(f"• Balance has decreased by {abs(balance_change):.1f}% over the period")
                
                for insight in cc_insights:
                    elements.append(Paragraph(insight, styles["Normal"]))
                elements.append(Spacer(1, 6))
            
            # Calculate and add insights for bank statements
            if bank_analyses:
                elements.append(Paragraph("Bank Account Insights:", styles["Heading4"]))
                bank_insights = []
                
                # Analyze bank account trends
                if len(bank_statement_data) > 1:
                    # Balance trend
                    first_balance = bank_statement_data[-1].get("Opening Balance", 0)
                    latest_balance = bank_statement_data[0].get("Closing Balance", 0)
                    balance_change = ((latest_balance - first_balance) / first_balance * 100) if first_balance != 0 else 0
                    if balance_change > 0:
                        bank_insights.append(f"• Account balance has increased by {balance_change:.1f}% over the period")
                    else:
                        bank_insights.append(f"• Account balance has decreased by {abs(balance_change):.1f}% over the period")
                    
                    # Deposit and withdrawal analysis
                    total_deposits = sum(data.get("Deposits", 0) for data in bank_statement_data)
                    total_withdrawals = sum(data.get("Withdrawals", 0) for data in bank_statement_data)
                    if total_deposits > 0 or total_withdrawals > 0:
                        bank_insights.append(f"• Total deposits: ${total_deposits:,.2f}")
                        bank_insights.append(f"• Total withdrawals: ${total_withdrawals:,.2f}")
                        if total_deposits > total_withdrawals:
                            bank_insights.append("• Net positive cash flow from deposits")
                        else:
                            bank_insights.append("• Net negative cash flow from withdrawals")
                
                for insight in bank_insights:
                    elements.append(Paragraph(insight, styles["Normal"]))
                elements.append(Spacer(1, 12))
            
            # Calculate trends
            total_purchases = []
            total_payments = []
            balances = []
            dates = []
            
            for analysis in analyses:
                if "output" in analysis and "analysis" in analysis["output"]:
                    data = analysis["output"]["analysis"]
                    if "balance_info" in data:
                        balance_info = data["balance_info"]
                        if all(key in balance_info for key in ["purchases", "payments", "new_balance"]):
                            total_purchases.append(float(balance_info["purchases"]))
                            total_payments.append(float(balance_info["payments"]))
                            balances.append(float(balance_info["new_balance"]))
                            if "statement_period" in data:
                                dates.append(data["statement_period"].get("end_date", "Unknown"))
            
            # Calculate trends only if we have valid data
            if len(total_purchases) > 1:
                # Calculate month-over-month percentage changes
                purchase_changes = []
                payment_changes = []
                balance_changes = []
                
                for i in range(len(total_purchases)-1):
                    # Calculate percentage change: (current - previous) / previous * 100
                    if total_purchases[i+1] != 0:
                        purchase_change = ((total_purchases[i] - total_purchases[i+1]) / total_purchases[i+1]) * 100
                        purchase_changes.append(purchase_change)
                    
                    if total_payments[i+1] != 0:
                        payment_change = ((total_payments[i] - total_payments[i+1]) / total_payments[i+1]) * 100
                        payment_changes.append(payment_change)
                    
                    if balances[i+1] != 0:
                        balance_change = ((balances[i] - balances[i+1]) / balances[i+1]) * 100
                        balance_changes.append(balance_change)
                
                # Calculate average trends
                purchase_trend = sum(purchase_changes) / len(purchase_changes) if purchase_changes else 0
                payment_trend = sum(payment_changes) / len(payment_changes) if payment_changes else 0
                balance_trend = sum(balance_changes) / len(balance_changes) if balance_changes else 0
                
                elements.append(Paragraph("Trend Analysis:", styles["Heading3"]))
                elements.append(Paragraph(f"• Purchase Trend: {'↑' if purchase_trend > 0 else '↓'} {abs(purchase_trend):.1f}%", styles["Normal"]))
                elements.append(Paragraph(f"• Payment Trend: {'↑' if payment_trend > 0 else '↓'} {abs(payment_trend):.1f}%", styles["Normal"]))
                elements.append(Paragraph(f"• Balance Trend: {'↑' if balance_trend > 0 else '↓'} {abs(balance_trend):.1f}%", styles["Normal"]))
                elements.append(Spacer(1, 12))
            
            # Add summary statistics
            elements.append(Paragraph("Summary Statistics:", styles["Heading3"]))
            for analysis in analyses:
                if "output" in analysis and "analysis" in analysis["output"]:
                    data = analysis["output"]["analysis"]
                    if "summary" in data:
                        summary = data["summary"]
                        elements.append(Paragraph(f"• Total Income: ${summary.get('total_income', 0):,.2f}", styles["Normal"]))
                        elements.append(Paragraph(f"• Total Expenses: ${summary.get('total_expenses', 0):,.2f}", styles["Normal"]))
                        elements.append(Paragraph(f"• Net Cash Flow: ${summary.get('net_cash_flow', 0):,.2f}", styles["Normal"]))
                        elements.append(Spacer(1, 6))
            
            # Category Analysis
            elements.append(Paragraph("Category Analysis", styles["Heading2"]))
            category_totals = {}
            
            for analysis in analyses:
                if "output" in analysis and "analysis" in analysis["output"]:
                    data = analysis["output"]["analysis"]
                    if "transactions" in data:
                        for trans in data["transactions"]:
                            category = trans.get("category", "other")
                            amount = trans.get("amount", 0)
                            category_totals[category] = category_totals.get(category, 0) + amount
            
            # Create category pie chart
            if category_totals:
                pie_chart = self._create_pie_chart(category_totals, "Expense Categories")
                elements.append(pie_chart)
                elements.append(Spacer(1, 12))
                
                # Add category breakdown
                elements.append(Paragraph("Category Breakdown:", styles["Heading3"]))
                for category, amount in sorted(category_totals.items(), key=lambda x: x[1], reverse=True):
                    elements.append(Paragraph(f"• {category.title()}: ${amount:,.2f}", styles["Normal"]))
                elements.append(Spacer(1, 12))
            
            # Process each analysis
            for analysis in analyses:
                if "output" in analysis and "analysis" in analysis["output"]:
                    data = analysis["output"]["analysis"]
                    
                    # Account Information
                    if "account_info" in data:
                        elements.append(Paragraph("Account Information", styles["Heading2"]))
                        account_info = data["account_info"]
                        elements.append(Paragraph(f"Account Number: {account_info.get('account_number', 'N/A')}", styles["Normal"]))
                        elements.append(Paragraph(f"Account Holder: {account_info.get('holder_name', 'N/A')}", styles["Normal"]))
                        elements.append(Spacer(1, 12))
                    
                    # Balance Information
                    if "balance_info" in data:
                        elements.append(Paragraph("Balance Information", styles["Heading2"]))
                        balance_info = data["balance_info"]
                        
                        # Create balance table data
                        balance_data = [["Item", "Amount"]]
                        
                        # Add credit card specific balance items
                        if "previous_balance" in balance_info:
                            balance_data.append(["Previous Balance", f"${balance_info['previous_balance']:,.2f}"])
                            balance_data.append(["Payments", f"-${balance_info['payments']:,.2f}"])
                            balance_data.append(["Other Credits", f"${balance_info['other_credits']:,.2f}"])
                            balance_data.append(["Purchases", f"${balance_info['purchases']:,.2f}"])
                            balance_data.append(["Cash Advances", f"${balance_info['cash_advances']:,.2f}"])
                            balance_data.append(["Fees", f"${balance_info['fees']:,.2f}"])
                            balance_data.append(["Interest", f"${balance_info['interest']:,.2f}"])
                            balance_data.append(["New Balance", f"${balance_info['new_balance']:,.2f}"])
                        else:
                            # Regular account balance items
                            balance_data.append(["Opening Balance", f"${balance_info['opening']:,.2f}"])
                            balance_data.append(["Closing Balance", f"${balance_info['closing']:,.2f}"])
                            balance_data.append(["Change", f"${balance_info['change']:,.2f}"])
                        
                        # Create and style the table
                        balance_table = Table(balance_data)
                        balance_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('FONTSIZE', (0, 0), (-1, 0), 14),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                            ('FONTSIZE', (0, 1), (-1, -1), 12),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ]))
                        elements.append(balance_table)
                        elements.append(Spacer(1, 12))
                    
                    # Recent Transactions
                    if "transactions" in data:
                        elements.append(Paragraph("Recent Transactions", styles["Heading2"]))
                        transactions = data["transactions"]
                        
                        if transactions:
                            # Sort transactions by date
                            def parse_date(date_str):
                                try:
                                    # Handle different date formats
                                    if ' ' in date_str:
                                        # Format: "Mar 11" or "Mar 11 NOCHES DE COLOMBIA..."
                                        date_str = date_str.split(' ')[0] + ' ' + date_str.split(' ')[1]
                                    return datetime.strptime(date_str, '%b %d')
                                except ValueError:
                                    return datetime.min

                            # Sort transactions by date
                            sorted_transactions = sorted(
                                transactions,
                                key=lambda x: parse_date(x['date']),
                                reverse=True  # Most recent first
                            )

                            # Create table data with sorted transactions
                            trans_data = [['Date', 'Description', 'Category', 'Amount', 'Type']]
                            for trans in sorted_transactions:
                                try:
                                    # Process description to prevent overlap
                                    desc = trans.get('description', 'N/A')
                                    if len(desc) > 30:
                                        desc = desc[:27] + '...'
                                    
                                    # Safely get date parts
                                    date_parts = trans.get('date', '').split(' ')
                                    date_str = ' '.join(date_parts[:2]) if len(date_parts) >= 2 else 'N/A'
                                    
                                    trans_data.append([
                                        date_str,
                                        desc,
                                        trans.get('category', 'other'),
                                        self._format_currency(trans.get('amount', 0)),
                                        trans.get('type', 'expense')
                                    ])
                                except Exception as e:
                                    print(f"Error processing transaction: {str(e)}")
                                    continue
                            
                            # Create and style the table with adjusted column widths
                            trans_table = Table(trans_data, colWidths=[60, 200, 80, 60, 60])
                            trans_table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Left align all cells
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 12),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                                ('FONTSIZE', (0, 1), (-1, -1), 10),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                                ('LEFTPADDING', (0, 0), (-1, -1), 6),  # Add left padding
                                ('RIGHTPADDING', (0, 0), (-1, -1), 6),  # Add right padding
                                ('TOPPADDING', (0, 0), (-1, -1), 6),  # Increase top padding
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),  # Increase bottom padding
                                ('WORDWRAP', (0, 0), (-1, -1), True),  # Enable word wrap for all cells
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Align text to top of cells
                                ('LEADING', (0, 0), (-1, -1), 12),  # Add line spacing for wrapped text
                                ('SPLITLONGWORDS', (0, 0), (-1, -1), True),  # Split long words if needed
                                ('SPLITROWS', (0, 0), (-1, -1), True),  # Allow rows to split across pages
                                ('NOSPLIT', (0, 0), (-1, 0)),  # Don't split header row
                                ('MINIMUMHEIGHT', (0, 0), (-1, -1), 30),  # Set minimum row height
                                ('TOPPADDING', (1, 1), (1, -1), 12),  # Extra padding for description column
                                ('BOTTOMPADDING', (1, 1), (1, -1), 12),  # Extra padding for description column
                                ('ALIGN', (1, 1), (1, -1), 'LEFT'),  # Left align description column
                                ('ALIGN', (0, 1), (0, -1), 'LEFT'),  # Left align date column
                                ('ALIGN', (2, 1), (2, -1), 'LEFT'),  # Left align category column
                                ('ALIGN', (3, 1), (3, -1), 'RIGHT'),  # Right align amount column
                                ('ALIGN', (4, 1), (4, -1), 'LEFT'),  # Left align type column
                                ('RIGHTPADDING', (1, 0), (1, -1), 12),  # Extra right padding for description column
                                ('LEFTPADDING', (2, 0), (2, -1), 12),  # Extra left padding for category column
                            ]))
                            elements.append(trans_table)
                            elements.append(Spacer(1, 12))
            
            # Build the PDF
            doc.build(elements)
            print(f"PDF report generated successfully at: {output_file}")
            
        except Exception as e:
            print(f"Error generating PDF report: {str(e)}")
            raise

class ReportGeneratorAgent:
    def __init__(self):
        self.llm = ChatOpenAI(
            temperature=0,
            model="gpt-4-turbo-preview",
            api_key=os.getenv("OPENAI_API_KEY")
        )
        
        self.prompt = ChatPromptTemplate.from_messages([
            SystemMessage(content=(
                "You are a financial report generator. Your task is to generate professional PDF reports "
                "from financial analysis results. The report should include:\n"
                "1. A clear summary of the financial data\n"
                "2. Visual representations of income and expenses\n"
                "3. Key insights and trends\n"
                "4. Detailed transaction information\n"
                "Use the report_generator tool to create well-formatted PDF reports."
            )),
            HumanMessage(content="Please generate a report from the following analysis data:\n\n{input}")
        ])
        
        self.tool = ReportGeneratorTool()
        
    def generate_report(self, analyses: List[Dict], output_file: str) -> Dict:
        """Generate a PDF report from financial analysis results."""
        try:
            # Generate the report directly using the tool
            self.tool._run(analyses, output_file)
            
            # Let the agent provide additional insights
            agent_result = self.llm.invoke(
                self.prompt.format_messages(input=str(analyses))
            )
            
            return {
                "status": "success",
                "data": {
                    "insights": agent_result.content
                }
            }
        except Exception as e:
            return {
                "status": "error",
                "error": str(e)
            }

    def generate_docx_report(self, analyses: List[Dict], output_file: str) -> Dict:
        """Generate a Word document report from financial analysis results."""
        try:
            # Generate the report directly using the tool
            self.tool._run(analyses, output_file)
            
            # Let the agent provide additional insights
            agent_result = self.llm.invoke(
                self.prompt.format_messages(input=str(analyses))
            )
            
            # Create a new Word document
            doc = Document()
            
            # Add Account Information section
            doc.add_heading('Account Information', level=1)
            
            account_info = analyses[0]['output']['analysis']['account_info']
            doc.add_paragraph(f"Account Number: {account_info.get('account_number', 'N/A')}")
            doc.add_paragraph(f"Account Holder: {account_info.get('holder_name', 'N/A')}")
            
            # Add Statement Period section
            doc.add_heading('Statement Period', level=1)
            
            period = analyses[0]['output']['analysis']['statement_period']
            doc.add_paragraph(f"From: {period.get('start_date', 'N/A')}")
            doc.add_paragraph(f"To: {period.get('end_date', 'N/A')}")
            
            # Add Balance Information section
            doc.add_heading('Balance Information', level=1)
            
            balance_info = analyses[0]['output']['analysis']['balance_info']
            doc.add_paragraph(f"Previous Balance: {self._format_currency(balance_info['previous_balance'])}")
            doc.add_paragraph(f"Payments: {self._format_currency(balance_info['payments'])}")
            doc.add_paragraph(f"Other Credits: {self._format_currency(balance_info['other_credits'])}")
            doc.add_paragraph(f"Purchases: {self._format_currency(balance_info['purchases'])}")
            doc.add_paragraph(f"Cash Advances: {self._format_currency(balance_info['cash_advances'])}")
            doc.add_paragraph(f"Fees: {self._format_currency(balance_info['fees'])}")
            doc.add_paragraph(f"Interest: {self._format_currency(balance_info['interest'])}")
            doc.add_paragraph(f"New Balance: {self._format_currency(balance_info['new_balance'])}")
            
            # Add Financial Summary section
            doc.add_heading('Financial Summary', level=1)
            
            summary = analyses[0]['output']['analysis']['summary']
            doc.add_paragraph(f"Total Income: {self._format_currency(summary['total_income'])}")
            doc.add_paragraph(f"Total Expenses: {self._format_currency(summary['total_expenses'])}")
            doc.add_paragraph(f"Net Cash Flow: {self._format_currency(summary['net_cash_flow'])}")
            
            # Add Recent Transactions section
            doc.add_heading('Recent Transactions', level=1)
            
            transactions = analyses[0]['output']['analysis']['transactions']
            if transactions:
                # Create transactions table with better formatting
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Set column widths
                table.columns[0].width = Inches(1.0)  # Date
                table.columns[1].width = Inches(3.0)  # Description
                table.columns[2].width = Inches(1.0)  # Category
                table.columns[3].width = Inches(1.0)  # Amount
                table.columns[4].width = Inches(1.0)  # Type
                
                # Add headers
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Date'
                header_cells[1].text = 'Description'
                header_cells[2].text = 'Category'
                header_cells[3].text = 'Amount'
                header_cells[4].text = 'Type'
                
                # Format headers
                for cell in header_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add transactions
                for transaction in transactions:
                    row_cells = table.add_row().cells
                    row_cells[0].text = transaction.get('date', '')
                    row_cells[1].text = transaction.get('description', '')
                    row_cells[2].text = transaction.get('category', '')
                    row_cells[3].text = self._format_currency(transaction.get('amount', 0))
                    row_cells[4].text = transaction.get('type', '')
                    
                    # Format cells
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        # Add word wrap for description
                        if cell == row_cells[1]:
                            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
                            cell.paragraphs[0].paragraph_format.widow_control = True
                            cell.paragraphs[0].paragraph_format.keep_with_next = False
            else:
                doc.add_paragraph("No transactions found.")
            
            return {
                "status": "success",
                "data": {
                    "insights": agent_result.content
                }
            }
        except Exception as e:
            return {
                "status": "error",
                "error": str(e)
            } 